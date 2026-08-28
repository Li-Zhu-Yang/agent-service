"""文档解析器：PDF / Word / Markdown / TXT → 纯文本。

统一清洗：去多余空行、去首尾空白。
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from core.exceptions import AppError

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"md", "txt", "pdf", "docx"}


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 合并连续空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_file(path: str | Path, file_type: str | None = None) -> str:
    """解析文件为纯文本。file_type 为 md/txt/pdf/docx 或带扩展名文件名。"""
    p = Path(path)
    ext = (file_type or p.suffix.lstrip(".")).lower().strip()
    if ext not in SUPPORTED_TYPES:
        raise AppError(f"暂不支持的文件类型: {ext}（支持 {sorted(SUPPORTED_TYPES)}）")

    raw = p.read_bytes()

    if ext == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(p)
            pages = [page.extract_text() or "" for page in reader.pages]
            return _clean("\n".join(pages))
        except Exception as exc:
            raise AppError(f"PDF 解析失败: {exc}") from exc

    if ext == "docx":
        try:
            from docx import Document

            doc = Document(p)
            parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return _clean("\n".join(parts))
        except Exception as exc:
            raise AppError(f"Word 解析失败: {exc}") from exc

    # md / txt
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            return _clean(raw.decode(enc))
        except UnicodeDecodeError:
            continue
    return _clean(raw.decode("utf-8", errors="replace"))


def parse_text(text: str) -> str:
    """解析已是文本的内容（同样清洗）。"""
    return _clean(text)
